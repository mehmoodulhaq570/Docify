import pandas as pd
from docx2pdf import convert
from pdf2docx import Converter
import pdfplumber
from docx import Document
import tempfile
from copy import deepcopy
import logging
import os


def word_to_pdf(input_file: str, output_file: str) -> None:
    try:
        if not input_file.lower().endswith(".docx"):
            raise ValueError("Input file must be a .docx file")
        convert(input_file, output_file)
        logging.info(f"Converted Word to PDF: {input_file} -> {output_file}")
    except Exception as e:
        logging.error(f"Error converting Word to PDF: {e}")
        raise


def pdf_to_word(
    input_file: str,
    output_file: str,
    preserve_images: bool = True,
    preserve_tables: bool = True,
    prefer_word: bool = False,
) -> None:
    try:
        if not input_file.lower().endswith(".pdf"):
            raise ValueError("Input file must be a .pdf file")
        # If user prefers Microsoft Word (Windows) and pywin32 is available, try it first
        if prefer_word:
            try:
                import sys

                if sys.platform.startswith("win"):
                    try:
                        import win32com.client

                        logging.info(
                            "Trying Microsoft Word COM conversion (preferred)..."
                        )
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc = word.Documents.Open(os.path.abspath(input_file))
                        # 16 = wdFormatDocumentDefault (docx)
                        doc.SaveAs(os.path.abspath(output_file), FileFormat=16)
                        doc.Close(False)
                        word.Quit()
                        logging.info(
                            f"Microsoft Word conversion succeeded (preferred): {input_file} -> {output_file}"
                        )
                        return
                    except Exception as word_exc:
                        logging.warning(
                            f"Preferred Microsoft Word conversion failed: {word_exc}"
                        )
            except Exception:
                # any import/platform error - ignore and continue
                pass
        cv = Converter(input_file)
        # Try primary conversion using pdf2docx
        try:
            cv.convert(output_file, start=0, end=None)
            cv.close()
            logging.info(f"Converted PDF to Word: {input_file} -> {output_file}")
            return
        except Exception as primary_exc:
            logging.warning(
                f"Primary pdf2docx conversion failed, attempting per-page conversion: {primary_exc}"
            )
            try:
                cv.close()
            except Exception:
                pass

            # Try per-page conversion and merge to preserve layout where possible.
            try:
                with pdfplumber.open(input_file) as pdf:
                    total_pages = len(pdf.pages)

                main_doc = Document()
                for page_idx in range(total_pages):
                    tmp = None
                    try:
                        # create a temp file for this single page conversion
                        tmpf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                        tmp = tmpf.name
                        tmpf.close()

                        cpage = Converter(input_file)
                        cpage.convert(tmp, start=page_idx, end=page_idx)
                        cpage.close()

                        # append converted page docx content into main_doc
                        tmp_doc = Document(tmp)
                        for element in tmp_doc.element.body:
                            main_doc.element.body.append(deepcopy(element))
                    except Exception as page_exc:
                        logging.warning(
                            f"Page {page_idx} conversion failed: {page_exc}; extracting text for this page instead."
                        )
                        # fallback for this page: extract text and add as paragraphs
                        try:
                            with pdfplumber.open(input_file) as pdf:
                                page = pdf.pages[page_idx]
                                text = page.extract_text()
                                if text:
                                    for line in text.split("\n"):
                                        main_doc.add_paragraph(line)
                        except Exception as text_exc:
                            logging.error(
                                f"Failed to extract text for page {page_idx}: {text_exc}"
                            )
                    finally:
                        if tmp:
                            try:
                                os.unlink(tmp)
                            except Exception:
                                pass

                # Save merged document
                main_doc.save(output_file)
                logging.info(
                    f"Per-page PDF->Word merge completed: {input_file} -> {output_file}"
                )
                return
            except Exception as per_page_exc:
                logging.warning(
                    f"Per-page conversion failed, falling back to text-only extraction: {per_page_exc}"
                )

            # Try Microsoft Word (Windows) conversion as a higher-fidelity fallback.
            # This requires Word to be installed and `pywin32` available in the Python env.
            try:
                import sys

                if sys.platform.startswith("win"):
                    try:
                        import win32com.client

                        logging.info(
                            "Trying Microsoft Word COM conversion as fallback..."
                        )
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        # Word can open a PDF and save it as a .docx (Word 2013+)
                        doc = word.Documents.Open(os.path.abspath(input_file))
                        # 16 = wdFormatDocumentDefault (docx)
                        doc.SaveAs(os.path.abspath(output_file), FileFormat=16)
                        doc.Close(False)
                        word.Quit()
                        logging.info(
                            f"Microsoft Word conversion succeeded: {input_file} -> {output_file}"
                        )
                        return
                    except Exception as word_exc:
                        logging.warning(
                            f"Microsoft Word conversion failed or not available: {word_exc}"
                        )
            except Exception:
                # any import/platform error - ignore and continue to text fallback
                pass

            # Final fallback: extract text and images using pdfplumber and write to a .docx using python-docx
            doc = Document()
            with pdfplumber.open(input_file) as pdf:
                for i, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        for line in text.split("\n"):
                            doc.add_paragraph(line)

                    # Extract and insert images from this page
                    try:
                        if hasattr(page, "images") and page.images:
                            for img_idx, img_info in enumerate(page.images):
                                try:
                                    # pdfplumber provides image info; we need to extract the image bytes
                                    # This is a best-effort extraction using the underlying page object
                                    # Note: pdfplumber doesn't directly extract image bytes easily, so we use a workaround
                                    # Use the page's underlying pdfminer object if available
                                    if hasattr(page, "page_obj") and hasattr(
                                        page.page_obj, "images"
                                    ):
                                        # Try to get image using page_obj (pdfminer.six structure)
                                        # This is a simplified attempt; real extraction can be complex
                                        pass  # Complex image extraction - skip for now to avoid errors

                                    # Alternative: try using img_info to locate the image stream
                                    # For now, we log that images exist but skip complex extraction
                                    logging.info(
                                        f"Image detected on page {i} but not extracted (complex)"
                                    )
                                except Exception as img_exc:
                                    logging.warning(
                                        f"Failed to extract image {img_idx} from page {i}: {img_exc}"
                                    )
                    except Exception as img_extract_exc:
                        logging.warning(
                            f"Image extraction failed for page {i}: {img_extract_exc}"
                        )

                    # don't add an extra page break after the last page
                    if i != len(pdf.pages) - 1:
                        doc.add_page_break()
            doc.save(output_file)
            logging.info(
                f"Fallback PDF->Word (text + attempted image extraction) completed: {input_file} -> {output_file}"
            )
            return
    except Exception as e:
        logging.error(f"Error converting PDF to Word: {e}")
        raise


def xlsx_to_csv(input_file: str, output_file: str) -> None:
    try:
        if not input_file.lower().endswith(".xlsx"):
            raise ValueError("Input file must be a .xlsx file")
        df = pd.read_excel(input_file)
        df.to_csv(output_file, index=False)
        logging.info(f"Converted Excel to CSV: {input_file} -> {output_file}")
    except Exception as e:
        logging.error(f"Error converting Excel to CSV: {e}")
        raise


def csv_to_xlsx(input_file: str, output_file: str) -> None:
    try:
        if not input_file.lower().endswith(".csv"):
            raise ValueError("Input file must be a .csv file")
        df = pd.read_csv(input_file)
        df.to_excel(output_file, index=False)
        logging.info(f"Converted CSV to Excel: {input_file} -> {output_file}")
    except Exception as e:
        logging.error(f"Error converting CSV to Excel: {e}")
        raise
